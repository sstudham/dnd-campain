from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


OUT = Path(r"C:\Users\studh\OneDrive\Documents\dnd-campain\output\docx\Session_14_DM_Notes.docx")

NAVY = "17324D"
BLUE = "2E5F85"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF4D6"
GOLD = "A66A12"
INK = "1F2933"
MUTED = "59636E"
WHITE = "FFFFFF"


def set_font(run, size=10.5, bold=False, italic=False, color=INK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_left_border(paragraph, color=BLUE, size=16, space=6):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, size=8.5, color=MUTED)


def add_numbering(doc, fmt, marker, left_twips=576, hanging_twips=288):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left_twips))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    ind.set(qn("w:hanging"), str(hanging_twips))
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.58)
section.bottom_margin = Inches(0.58)
section.left_margin = Inches(0.68)
section.right_margin = Inches(0.68)
section.header_distance = Inches(0.28)
section.footer_distance = Inches(0.28)

# Compact tabletop-print override to compact_reference_guide.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(2)
normal.paragraph_format.line_spacing = 1.0

for style_name, size, color, before, after in [
    ("Heading 1", 14, NAVY, 8, 3),
    ("Heading 2", 11.5, BLUE, 6, 2),
    ("Heading 3", 10.5, GOLD, 5, 1.5),
]:
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.line_spacing = 1.0

read_style = doc.styles.add_style("Read Aloud", WD_STYLE_TYPE.PARAGRAPH)
read_style.font.name = "Calibri"
read_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
read_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
read_style.font.size = Pt(9.7)
read_style.font.italic = True
read_style.font.color.rgb = RGBColor.from_string(INK)
read_style.paragraph_format.left_indent = Inches(0.18)
read_style.paragraph_format.right_indent = Inches(0.12)
read_style.paragraph_format.space_before = Pt(0)
read_style.paragraph_format.space_after = Pt(2)
read_style.paragraph_format.line_spacing = 1.0

bullet_num = add_numbering(doc, "bullet", "•")
decimal_num = add_numbering(doc, "decimal", "%1.")


def p(text="", bold_lead=None, italic=False, color=INK, after=None, keep=False):
    para = doc.add_paragraph()
    para.paragraph_format.keep_together = keep
    if after is not None:
        para.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        first = para.add_run(bold_lead)
        set_font(first, bold=True, color=color)
        rest = para.add_run(text[len(bold_lead):])
        set_font(rest, italic=italic, color=color)
    else:
        run = para.add_run(text)
        set_font(run, italic=italic, color=color)
    return para


def h1(text):
    return doc.add_paragraph(text, style="Heading 1")


def h2(text):
    return doc.add_paragraph(text, style="Heading 2")


def h3(text):
    return doc.add_paragraph(text, style="Heading 3")


def bullet(text, bold_lead=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(1.5)
    para.paragraph_format.line_spacing = 1.0
    apply_num(para, bullet_num)
    if bold_lead and text.startswith(bold_lead):
        r1 = para.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = para.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        set_font(para.add_run(text))
    return para


def number(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(1.5)
    para.paragraph_format.line_spacing = 1.0
    apply_num(para, decimal_num)
    set_font(para.add_run(text))
    return para


def callout(label, text, fill=PALE_BLUE, border=BLUE):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.14)
    para.paragraph_format.right_indent = Inches(0.12)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.0
    shade_paragraph(para, fill)
    set_left_border(para, border)
    r1 = para.add_run(label + " ")
    set_font(r1, bold=True, color=NAVY)
    r2 = para.add_run(text)
    set_font(r2)
    return para


def read(text):
    para = doc.add_paragraph(style="Read Aloud")
    shade_paragraph(para, "F4F7F9")
    set_left_border(para, GOLD, size=12, space=5)
    set_font(para.add_run(text), size=9.7, italic=True)
    return para


def page_break():
    doc.add_page_break()


# Header and footer
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.paragraph_format.space_after = Pt(0)
set_font(header_p.add_run("TUG4  |  SESSION 14  |  CLIFFTOP OBSERVATORY"), size=8, bold=True, color=MUTED)

footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_p.paragraph_format.space_before = Pt(0)
footer_p.paragraph_format.space_after = Pt(0)
set_font(footer_p.add_run("DM NOTES  •  PAGE "), size=8, color=MUTED)
add_page_field(footer_p)


# PAGE 1
kicker = doc.add_paragraph()
kicker.paragraph_format.space_before = Pt(0)
kicker.paragraph_format.space_after = Pt(0)
set_font(kicker.add_run("TONIGHT'S RUN SHEET"), size=9, bold=True, color=GOLD)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(1)
set_font(title.add_run("SESSION 14 DM NOTES"), size=22, bold=True, color=NAVY)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_before = Pt(0)
subtitle.paragraph_format.space_after = Pt(5)
set_font(subtitle.add_run("Clifftop Observatory: Rescue Aidron"), size=12.5, bold=True, color=BLUE)

callout(
    "RUN TONIGHT IN ONE SENTENCE:",
    "Sparkrender plans to sacrifice Aidron and steal the power of Stormwreck's dead dragons; the party must learn the tower's secret, rescue Aidron, and decide when to confront the blue dragon.",
)

h2("The Situation")
p("Sparkrender is a young blue dragon occupying the Clifftop Observatory. He captured Aidron, a young bronze dragon from Dragon's Rest, and sealed him beneath the main tower. The Observatory's ancient instrument is aligning the power left by dragons that died on Stormwreck Isle. Sparkrender intends to use Aidron as the living key in a ritual that will bind that power to him.")

h2("Expected Flow")
for item in [
    "Recover from the stirge fight and speak with Minn.",
    "Handle Sparkrender's kobolds in D3 without assuming combat.",
    "Explore D4, recover the statue clue, and find the lightning-resistance potion.",
    "Enter D5 and choose whether to wake Sparkrender or sneak past him.",
    "Open the hidden descent to D6 and rescue Aidron.",
    "Confront Sparkrender with Aidron's help, tonight or next session.",
]:
    number(item)

callout(
    "LIKELY STOP:",
    "Aidron is freed and the party commits to a plan. Stretch goal: return to D5 and roll initiative. Do not rush the final battle just to finish the published adventure tonight.",
    fill=PALE_GOLD,
    border=GOLD,
)

h2("Opening Narration - Read Aloud")
for text in [
    "Last time, Runara sent you toward the Clifftop Observatory. The sickness beneath Stormwreck had shown itself in Seagrow Caves. The dead around the island had shown themselves aboard the Compass Rose. Whatever was gathering above the island was waiting among the ruined towers.",
    "The moonstone key awakened two marble dragons, and a bridge of light reached across empty air. For Floyd, the cliff became Glittervein again: darkness, falling stone, the silence after the last pickaxe stopped, and the golden ray that found him when no one else could.",
    'A voice told him, "You were not spared to find gold. You were spared to find the way out."',
    "Across the bridge, eight stirges descended on the ruined rotunda, where Mek and Minn were already fighting for their lives. Every one of you killed at least one. Mek did not survive. Minn did.",
    'Severed Whisper drove his blade through the final creature and said, "I would like to introduce myself. I am sorry..."',
    "Then the five dragon effigies began to make sense: bronze, gold, brass, blue, and red. The blue light was not studying the others. It was forcing them into alignment.",
    'Floyd heard his fallen crew and the rhythm of their pickaxes. His holy symbol burned gold. "You were not the only one buried. Dig."',
    "Now the stirges lie dead across the Observatory floor. Mek lies still beside the instrument. Minn kneels over his brother. The blue light turns once around the golden rings.",
    "From somewhere inside the tower comes a heavy crash. Dust spills from the stones beneath your feet. Something is alive down there. Somewhere above it, thunder is breathing.",
    "What do you do?",
]:
    read(text)


# PAGE 2
page_break()
h1("1. Recover, Question Minn, and Reach D4")

h2("Party Condition and Immediate Fixes")
p("The published Observatory normally expects Level 3 characters. TUG4 has five Level 2 characters, and four are wounded. Allow one safe short rest in D2. A short rest does not trigger the ritual; leaving or attempting a long rest will.")
for item in [
    "Brother Kai: 4/13 HP, AC 14.",
    "Pat Benatar: 5/13 HP, AC 13.",
    "Severed Whisper: 13/13 HP, AC 17.",
    "Throk: 4/13 HP, AC 12. He has the only listed Potion of Healing.",
    "Floyd GoldSeeker: 6/13 HP, AC 13. Divine Spark can heal 1d8 per Channel Divinity use.",
]:
    bullet(item)

callout(
    "BEFORE MOVING ON:",
    "Let everyone spend Hit Dice. Ask Andy to confirm Floyd's spells and slots. Floyd should equip his chain shirt with his shield for AC 16. Throk can remove leather and use Armor of Shadows for AC 14.",
    fill=PALE_GOLD,
    border=GOLD,
)

h2("D2 Rotunda Aftermath")
p("Minn kneels beside Mek, repeatedly pressing both hands against his brother's chest as if the correct pressure might restart him. The golden instrument continues turning. Blue energy pulses through it. Bronze light flickers downward toward the base of the tallest tower. Another heavy impact comes from beneath that tower.")

h3("Minn Knows")
for item in [
    "Sparkrender ordered the kobolds to create and arrange the five dragon effigies.",
    "He is waiting for the instrument to show that the ritual's timing is right.",
    "A bronze-colored dragon arrived, argued with Sparkrender, and was trapped below the tower.",
    "Sparkrender ordered the kobolds not to touch the collapsed wall at the tower's base.",
    "Other kobolds are camped in D3; an abandoned study stands across the gap at D4.",
]:
    bullet(item)

h3("Useful Minn Lines")
for item in [
    '"Mek said the blue one would make us important."',
    '"The other dragon is still making noise. That means he is still alive."',
    '"Sparkrender said not to move the fallen wall. He said the stone was doing its job."',
    '"If Myla asks, I do not know what I am supposed to tell her."',
]:
    bullet(item)
p("Do not make Minn blame the party. Sparkrender and the Observatory's dangers are responsible for Mek's death.")

h2("D3 Kobold Camp")
p("Ekrash, Erp, Hev, Nuhro, and Snirke occupy the camp; Nuhro and Snirke have wings. They begin defensive and threatening, but this is a social problem unless the players choose violence.")
for item in [
    "If Minn speaks first, shift the camp from hostile to frightened and suspicious.",
    "Ekrash is the loud loyalist. Nuhro watches for an escape. Snirke wants permission to survive. Erp and Hev follow the strongest voice.",
    "Use DC 13 for Intimidation or DC 15 for Deception/Persuasion. Lower or skip the roll when Minn supports the party or the players address the kobolds' fear.",
    "They can reveal that Sparkrender sleeps in D5, Aidron is below him, and D4 contains old books Sparkrender could not open.",
    "A winged kobold can retrieve the moonstone key if it remained in the D1 bridge anchor.",
]:
    bullet(item)
p("A good outcome is not recruitment. The kobolds may stand aside, fetch the key, keep watch, or promise not to warn Sparkrender.")

h2("D4 Isolated Study")
p("The study lies across a 22-foot gap. The clean route uses the moonstone key in D2 to create a second energy bridge. Reward workable alternatives involving ropes, climbing, jumping, or swimming.")

read("The bridge leads to a cramped tower open to the wind. Waves strike the rocks far below through a hole where half the floor has collapsed. Ruined books cover the floor. One small black journal remains intact beneath a slab of marble. Its lock is perfectly clean.")

h3("Journal Trap and Clue")
for item in [
    "DC 15 Perception notices the arcane rune.",
    "DC 11 Arcana explains how to disable it by scratching through the rune.",
    "DC 10 Dexterity with thieves' tools opens the lock; DC 12 Strength breaks it.",
    "Opening it without disabling the rune deals 1d6 poison damage to the opener.",
]:
    bullet(item)
callout(
    "THE CLUE:",
    '"Four scholars turn their eyes toward the Dragon of Dawn. Where their sight meets, the descent into hidden knowledge begins."',
)
p("The journal shows the Observatory's purpose changing over time: first listening to the island's dragon energy, then drawing it, then taking it. Do not explain the larger rift. If Pat handles the journal, her clan-marked scale briefly warms without explanation.")
p("Floyd notices a loose brick because its weight sits incorrectly. Behind it are 10 gp and a Potion of Lightning Resistance. Make sure the party has a fair chance to find this protection.")


# PAGE 3
page_break()
h1("2. D5 - The Sleeping Dragon and Hidden Stair")

read("Broken stained glass scatters colored light over a star map set into the floor with gold and tiny stones. Four enormous scholar statues stand around the room, each facing a different direction.")
read("Sparkrender sleeps in the northeast corner among coins and blue gems. Lightning crawls between his horns with every breath.")
read("Directly beneath him, something strikes the tower wall.")

h2("Make the Choice Explicit")
for item in [
    "Attack Sparkrender now.",
    "Sneak around him and solve the statue puzzle.",
    "Climb down outside and excavate the collapsed wall.",
    "Try talking to him.",
    "Retreat and make another plan.",
]:
    bullet(item)
p("Do not hide the rescue route behind a roll. Rolls determine noise, time, and cost - not whether the players understand their options.")

h2("The Statue Puzzle")
for item in [
    "DC 10 Investigation finds the dragon-shaped constellation in the southeast part of the star map.",
    "DC 15 Perception notices that the statues rotate.",
    "DC 10 Perception while turning a statue reveals when it settles into the correct position.",
    "When all four statues face the Dragon of Dawn, part of the floor descends into a spiral staircase to D6.",
]:
    bullet(item)

h2("Sneaking Past Sparkrender")
for item in [
    "At least half the party must succeed on DC 14 Stealth to move around him quietly.",
    "Turning the north and east statues closest to him requires separate DC 14 Stealth checks.",
    "The statues provide cover if combat begins.",
    "Lightning crawling across the floor telegraphs when his breath is charged.",
]:
    bullet(item)

h2("Talking to Sparkrender")
p("A Draconic speaker can delay immediate violence with a DC 12 Deception, Intimidation, or Persuasion check. Mentioning Runara or Aidron ends his patience.")

h3("How to Play Him")
for item in [
    "Vain, ambitious, and insecure about being young and small.",
    "He believes the power of dead dragons belongs to whoever is strong enough to take it.",
    'He calls the dead dragons "ore left in the ground."',
    'He calls Aidron "the living key."',
    "He considers Runara's restraint weakness.",
    "He misreads Floyd's survival as proof that survivors are superior.",
]:
    bullet(item)

h3("What He Might Offer")
p("Sparkrender may promise the party status beneath him if they help complete the ritual. He can boast long enough for clever players to study the room, move into position, or buy time. Do not force a long villain speech.")

h2("If Combat Starts in D5")
for item in [
    "AC 17; HP 52; immune to lightning.",
    "Bite: +5 to hit, piercing plus lightning damage.",
    "Lightning Breath: 30-foot line, DC 12 Dexterity save, 4d10 lightning damage, half on success; recharge 5-6.",
    "Use the statues as cover and keep the party from lining up.",
    "In the normal D5 encounter, Sparkrender tries to flee at 10 HP or fewer.",
]:
    bullet(item)
callout(
    "DANGER:",
    "At current HP, one breath can drop several characters. The fair route is short rest, D4 resistance potion, rescue Aidron, then fight. Do not add the D3 kobolds unless the party deliberately created an ongoing battle.",
    fill=PALE_GOLD,
    border=GOLD,
)


# PAGE 4
page_break()
h1("3. D6 - Rescue Aidron and Set the Finale")

h2("Entering the Secret Library")
p("The party can use the hidden staircase or clear the collapsed outside wall. Clearing rubble takes one person 30 minutes, two people 15 minutes, or the whole party about 6 minutes. Quiet work takes twice as long and requires at least half the party to succeed on DC 14 Stealth.")
p("Floyd can identify which stones are safe to move, but the rescue remains possible without him.")

read("The passage opens into stale air, old parchment, and broken glass. Shelves cover the walls, their books swollen with damp and age.")
read("A desk explodes into splinters. A bronze dragon no larger than a bear rises from the wreckage. Dust covers his scales. One wing hangs low. His claws are bloodied from digging.")
read('He looks at the open passage and asks, "Did Runara send you?" Above, thunder answers for you.')

h2("Aidron")
p("Aidron is proud, frightened, exhausted, and embarrassed that Sparkrender defeated him. He is not helpless, but he needs allies.")
for item in [
    "He confronted Sparkrender because he believed Runara was being too cautious.",
    "Sparkrender defeated him and sealed him inside the library.",
    "Sparkrender intends to kill Aidron during the ritual.",
    "Aidron's death will let Sparkrender seize the power of Stormwreck's dead dragons.",
    "Aidron wants to attack immediately, but he will hear a reasonable plan if treated as an ally.",
]:
    bullet(item)
p("The library contains either a +1 battleaxe or a Hold Person spell scroll. DC 15 Investigation, Detect Magic, or Aidron's help reveals it.")

h2("Running the Dragon Fight")
p("With Aidron helping, five Level 2 characters can drive off Sparkrender, but positioning matters more than damage.")
for item in [
    "Spread out before combat so the lightning line cannot hit several characters.",
    "Give the Potion of Lightning Resistance to a wounded front-line character.",
    "Use the statues as cover.",
    "Aidron is immune to lightning, but his bite and presence add a crucial extra turn.",
    "In D5, Sparkrender flees at 10 HP or fewer. If the ritual has already begun, he fights to the death.",
]:
    bullet(item)

h2("Choose the Ending That Fits the Clock")
h3("A. The Sleeping Dragon")
p("If the party reaches D5 late, end when they see Sparkrender sleeping over the star map and hear Aidron beneath him. Ask whether each character's first instinct is attack, rescue, bargain, or retreat.")
h3("B. The Dragon Beneath")
p('If they free Aidron late, blue light leaks through the ceiling. Aidron looks up and says, "He has started without me. That means he is afraid of you."')
h3("C. Roll Initiative")
p("If they move quickly, return to D5 with Aidron, wake Sparkrender, roll initiative, and end as lightning crosses the star map.")

h2("If the Party Leaves or Takes a Long Rest")
p("Sparkrender begins the ritual in D2. Aidron is restrained by three heavy chains. One action opens one clasp; all three must be released. In this version Sparkrender fights to the death because the ritual is underway.")

h2("Table Checklist")
for item in [
    "Screen image: Session 13 Final gathering at the magical temple.png.",
    "Observatory map showing D2-D6.",
    "Tokens: five PCs, Minn, five D3 kobolds, Sparkrender, and Aidron.",
    "D4 journal clue, Potion of Lightning Resistance, and 10 gp.",
    "Four directional markers for the D5 statue puzzle.",
    "A d6 for Sparkrender's breath recharge.",
]:
    bullet(item)

callout(
    "THE ONLY REQUIRED OUTCOME:",
    "The players understand that Aidron is trapped, choose how to reach him, and make a deliberate decision about Sparkrender. Completing the whole adventure tonight is optional.",
)


# Keep headings and short callouts together where possible.
for para in doc.paragraphs:
    if para.style.name.startswith("Heading"):
        para.paragraph_format.keep_with_next = True
    if para.style.name in {"Heading 2", "Heading 3"}:
        para.paragraph_format.widow_control = True

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
